import io
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.data.db import get_db
from datetime import datetime

from app.data.autoparte import Autoparte
from app.data.usuario_interno import UsuarioInterno
from app.data.usuario_externo import UsuarioExterno
from app.data.pedido import Pedido

# Librerías PDF
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# Librerías Excel
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

# Librerías Word
from docx import Document
from docx.shared import Pt, RGBColor

router = APIRouter(prefix="/v1/reportes", tags=['Módulo de Reportes'])

# --- MAGIA DE FILTRADO AQUÍ ---
def extraer_datos_estructurados(tipo: str, db: Session, filtro: str, fecha_inicio: Optional[str] = None, fecha_fin: Optional[str] = None):
    if tipo == "inventario":
        headers = ["CÓDIGO", "NOMBRE", "MARCA", "CATEGORÍA", "STOCK"]
        query = db.query(Autoparte)
        if filtro != "todos":
            query = query.filter(Autoparte.categoria == filtro) # Filtra por categoría
        filas = [[p.codigo, p.nombre, p.marca, p.categoria, str(p.stock)] for p in query.all()]
        return headers, filas

    elif tipo == "empleados":
        headers = ["NOMBRE", "CORREO", "DEPARTAMENTO", "ESTATUS"]
        query = db.query(UsuarioInterno)
        if filtro == "activos":
            query = query.filter(UsuarioInterno.activo == True)
        elif filtro == "inactivos":
            query = query.filter(UsuarioInterno.activo == False)
        filas = [[e.nombre, e.correo, e.departamento, "Activo" if e.activo else "Baja"] for e in query.all()]
        return headers, filas

    elif tipo == "clientes":
        headers = ["NOMBRE", "CORREO", "TELÉFONO"]
        filas = [[c.nombre, c.correo, c.telefono] for c in db.query(UsuarioExterno).all()]
        return headers, filas

    elif tipo == "pedidos":
        headers = ["FOLIO", "ID CLIENTE", "FECHA", "ESTATUS"]
        query = db.query(Pedido)
        
        # Filtro 1: Estatus
        if filtro != "todos":
            query = query.filter(Pedido.estatus == filtro)
            
        # Filtro 2: Rango de Fechas (Si el usuario las mandó)
        if fecha_inicio:
            query = query.filter(Pedido.fecha >= fecha_inicio)
        if fecha_fin:
            query = query.filter(Pedido.fecha <= fecha_fin)
            
        filas = [[str(p.id), str(p.id_cliente), p.fecha.strftime('%Y-%m-%d'), p.estatus] for p in query.all()]
        return headers, filas
    return None, None

# Agregamos "filtro" como parámetro opcional (por defecto "todos")
@router.get("/{tipo}/{formato}")
async def generar_reporte_profesional(tipo: str, formato: str, filtro: str = "todos", fecha_inicio: Optional[str] = None, fecha_fin: Optional[str] = None, db: Session = Depends(get_db)):
    
    headers, filas = extraer_datos_estructurados(tipo, db, filtro, fecha_inicio, fecha_fin)
    
    if headers is None:
        raise HTTPException(status_code=404, detail="Reporte no válido.")
    if not filas:
        filas = [["No hay datos para estos filtros", "", "", "", ""][:len(headers)]]

    buffer = io.BytesIO()

    # --- 1. GENERADOR PDF (DISEÑO CORPORATIVO CORREGIDO) ---
    if formato == "pdf":
        # Usamos landscape para acostar la hoja y reducimos los márgenes
        doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
        elementos = []
        estilos = getSampleStyleSheet()
        
        # Título
        titulo = Paragraph(f"<b>REPORTE OFICIAL MACUIN: {tipo.upper()}</b>", estilos['Title'])
        elementos.append(titulo)
        elementos.append(Spacer(1, 20))
        
        # Cálculo Dinámico de Columnas (720 puntos aprox en hoja horizontal)
        datos_tabla = [headers] + filas
        ancho_disponible = 720 
        ancho_columna = ancho_disponible / len(headers)
        
        tabla = Table(datos_tabla, colWidths=[ancho_columna]*len(headers))
        
        # Estilos de la tabla (Letra tamaño 9 y márgenes internos)
        tabla.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#6B0F2A")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9), 
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#F7FAFC")),
            ('GRID', (0, 0), (-1, -1), 1, colors.silver),
        ]))
        elementos.append(tabla)
        doc.build(elementos)
        media_type = "application/pdf"

    # --- 2. GENERADOR EXCEL ---
    elif formato == "xlsx":
        wb = Workbook()
        ws = wb.active
        ws.title = f"Reporte {tipo}"
        
        ws.append(headers)
        guinda_fill = PatternFill(start_color="6B0F2A", end_color="6B0F2A", fill_type="solid")
        white_font = Font(color="FFFFFF", bold=True)
        
        for col, _ in enumerate(headers, start=1):
            celda = ws.cell(row=1, column=col)
            celda.fill = guinda_fill
            celda.font = white_font
            celda.alignment = Alignment(horizontal="center")
            ws.column_dimensions[celda.column_letter].width = 25 # Hicimos las celdas de Excel más anchas también
            
        for fila in filas:
            ws.append(fila)
            
        wb.save(buffer)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    # --- 3. GENERADOR WORD ---
    elif formato == "docx":
        doc = Document()
        titulo = doc.add_heading(f'Reporte de {tipo.capitalize()} - MACUIN', 0)
        titulo.runs[0].font.color.rgb = RGBColor(107, 15, 42)
        
        tabla = doc.add_table(rows=1, cols=len(headers))
        tabla.style = 'Light Grid Accent 1'
        
        hdr_cells = tabla.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            
        for fila in filas:
            row_cells = tabla.add_row().cells
            for i, valor in enumerate(fila):
                row_cells[i].text = str(valor)
                
        doc.save(buffer)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    buffer.seek(0)
    return StreamingResponse(buffer, media_type=media_type, headers={"Content-Disposition": f"attachment; filename=Reporte_Macuin_{tipo}.{formato}"})